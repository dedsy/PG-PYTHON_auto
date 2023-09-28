import openpyxl
from docxtpl import DocxTemplate
import re
from docx import Document


def generate_leave_application(excel_file: str, excel_sheet: str, doc_template: str, email_errors_file: str):
    """Функция проходится по Excel файлу и заполняет заявления на отпуск по шаблону, сохраняет в отдельный для каждого сотрудника файл.
    Так же проводится проверка на правильность введённого mail и запись его в виде таблицы в отдельный файл."""

    # Открывает файл с сотрудниками
    wb = openpyxl.load_workbook(filename=excel_file)
    sheet = wb[excel_sheet]

    # Вызываем класс DocxTemplate для формирования заявления на отпуск
    doc = DocxTemplate(doc_template)

    # Вызываем класс Document из модуля docx для формирования файла с ошибками в email
    email_doc = Document()
    # Создаём таблицу и заполняем названия колонок
    table = email_doc.add_table(rows=1, cols=2, style='Table Grid')
    table.cell(0, 0).paragraphs[0].add_run('ФИО').bold = True
    table.cell(0, 1).paragraphs[0].add_run('email').bold = True

    # Считаем колличество строк в Excel файле
    row_count = sheet.max_row

    # Создаём словарь для неправильных email
    wrong_emails = {}

    for i in list(range(2, row_count+1)):
        # Создаём переменные с данными из Excel файла
        last_name = sheet['A' + str(i)].value
        name = sheet['B' + str(i)].value
        company = sheet['D' + str(i)].value
        vacation_start = sheet['F' + str(i)].value.strftime("%d.%m.%Y")
        vacation_end = sheet['G' + str(i)].value.strftime("%d.%m.%Y")
        email = sheet['E' + str(i)].value

        # Создаём словарь для маппинга полей в шаблоне заявления
        content = {'last_name': last_name,
                   'name': name,
                   'company': company,
                   'vacation_start': vacation_start,
                   'vacation_end': vacation_end}

        # Применяем наш словарь на шаблон
        doc.render(content)
        # Сохраняем файл заявления
        doc.save(str(i)+'_'+last_name+'.docx')

        # Строка валидации email
        regex = re.compile(r'([A-Za-z0-9]+[.-_])*[A-Za-z0-9]+@[A-Za-z0-9-]+(\.[A-Z|a-z]{2,})+')

        # Создаём переменную содержащую фамилию и имя
        full_name = str(last_name + " " + name)

        # Записываем неправильные email в словарь
        if not re.fullmatch(regex, email):
            d = {full_name: email}
            wrong_emails.update(d)

    # Если словарь не пустой, добавляем строку и заполняем в файле неправильных email
    if len(wrong_emails) > 0:
        for k, v in wrong_emails.items():
            row = table.add_row()
            row.cells[0].text = k
            row.cells[1].text = v

    # Сохраняем файл неправильных email
    email_doc.save(email_errors_file)
